"""
Convert legal Markdown sources to professionally formatted A4 .docx files.

Targets:
  - docs/legal/contrat-services-saas.md      → contrat-services-saas.docx
  - docs/legal/fiche-souscription-client.md  → fiche-souscription-client.docx

Run with no arguments to build both, or pass a single file path:
  python md_to_docx.py                                     # build all
  python md_to_docx.py docs/legal/contrat-services-saas.md # build one

Design choices:
- A4 page size, Calibri 11pt body, 1 inch margins
- Smart quotes everywhere (apostrophes courbes, guillemets français « »)
- Heading 1 (#)  → CONTRAT title (centered, 18pt bold, page 1)
- Heading 2 (##) → ARTICLE / ANNEXE (14pt bold, blue, with top spacing)
- Heading 3 (###) → sub-sections inside annexes (12pt bold)
- Tables with grey header row, centered column widths
- HR (---) as a real bottom-border paragraph (skill rule: never use empty tables)
- Strips the trailing HTML <!-- ... --> editor-notes block before rendering
- Footer with page numbers
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ─── Paths ──────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parents[3]
LEGAL_DIR = ROOT / "docs" / "legal"
DEFAULT_SOURCES = [
    LEGAL_DIR / "contrat-services-saas.md",
    LEGAL_DIR / "fiche-souscription-client.md",
]


# ─── Smart-quote / typography pass ──────────────────────────────────────
def smart_quotes(text: str) -> str:
    """Replace straight quotes with French typography. Only on plain text — never on code."""
    # Apostrophes : ' → ’
    text = re.sub(r"(\w)'(\w)", r"\1’\2", text)
    text = text.replace(" '", " ’")
    text = text.replace("' ", "’ ")
    # Guillemets français « » sont déjà dans la source, on garde
    return text


# ─── Inline parsing ─────────────────────────────────────────────────────
INLINE_CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")


def add_inline(p, text: str, base_font="Calibri", base_size=11, base_bold=False):
    """Add a stretch of inline-formatted markdown into a paragraph."""
    text = smart_quotes(text)

    # Tokenize : alternance of code|bold|italic|plain
    pattern = re.compile(
        r"(`[^`]+`|\*\*[^*]+\*\*|(?<!\*)\*[^*]+\*(?!\*))"
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = p.add_run(text[pos:m.start()])
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.bold = base_bold
        token = m.group(0)
        if token.startswith("**"):
            run = p.add_run(token[2:-2])
            run.bold = True
            run.font.name = base_font
            run.font.size = Pt(base_size)
        elif token.startswith("`"):
            inner = token[1:-1]
            run = p.add_run(inner)
            run.font.name = "Consolas"
            run.font.size = Pt(base_size - 1)
            # Detect placeholder patterns: `[...]` or `…` (ellipsis) → yellow
            # highlight to make them impossible to miss when filling. Other
            # backticked text (technical strings) stays light-grey.
            is_placeholder = (
                inner.startswith("[") and inner.endswith("]")
            ) or inner.strip() in ("…", "...")
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            if is_placeholder:
                shd.set(qn("w:fill"), "FFE699")  # bright fillable yellow
                run.bold = True
                run.font.color.rgb = RGBColor(0x80, 0x4A, 0x00)
            else:
                shd.set(qn("w:fill"), "F2F2F2")
            run._element.get_or_add_rPr().append(shd)
        elif token.startswith("*"):
            run = p.add_run(token[1:-1])
            run.italic = True
            run.font.name = base_font
            run.font.size = Pt(base_size)
        pos = m.end()
    if pos < len(text):
        run = p.add_run(text[pos:])
        run.font.name = base_font
        run.font.size = Pt(base_size)
        run.bold = base_bold


# ─── Style helpers ──────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_horizontal_rule(doc):
    """Use a paragraph with a bottom border, NOT a table (skill rule)."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Run with field "PAGE / NUMPAGES"
    run = p.add_run()
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)

    p.add_run(" / ").font.size = Pt(9)

    run2 = p.add_run()
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    fld_begin2 = OxmlElement("w:fldChar")
    fld_begin2.set(qn("w:fldCharType"), "begin")
    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = " NUMPAGES "
    fld_end2 = OxmlElement("w:fldChar")
    fld_end2.set(qn("w:fldCharType"), "end")
    run2._r.append(fld_begin2)
    run2._r.append(instr2)
    run2._r.append(fld_end2)


# ─── Markdown parser (line-based, hand-rolled for our restricted dialect) ─
def is_table_row(line: str) -> bool:
    return line.startswith("|") and line.endswith("|")


def parse_table_block(lines, idx):
    """Return (rows: list[list[str]], next_idx)."""
    rows = []
    while idx < len(lines) and is_table_row(lines[idx]):
        cells = [c.strip() for c in lines[idx].strip("|").split("|")]
        rows.append(cells)
        idx += 1
    # rows[1] is the alignment row (---|---|---) — drop it
    if len(rows) >= 2 and all(re.match(r":?-+:?$", c) for c in rows[1]):
        rows = [rows[0]] + rows[2:]
    return rows, idx


def render_table(doc, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row):
            cell = tr.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Clear default paragraph
            cell.text = ""
            p = cell.paragraphs[0]
            if r_idx == 0:
                set_cell_shading(cell, "1F4E79")  # bleu profond pour entête
                add_inline(p, cell_text, base_size=10, base_bold=True)
                # White text on blue header
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                add_inline(p, cell_text, base_size=10)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)


def render_signature_block(doc):
    """Replace the markdown 'Signatures' section with a clean visual block:
    two side-by-side signature boxes with explicit e-signature anchor markers
    that DocuSign / Yousign / Adobe Sign can auto-detect."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline(
        p,
        "Fait en deux (2) exemplaires originaux, signés électroniquement, "
        "à `[Ville]`, le `[JJ/MM/AAAA]`. La signature électronique a la même "
        "valeur juridique qu'une signature manuscrite conformément à la loi "
        "marocaine n° 53-05 relative à l'échange électronique de données "
        "juridiques.",
    )

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cells = table.rows[0].cells
    headers = ["POUR LE PRESTATAIRE", "POUR LE CLIENT"]
    anchors = [
        ("PRESTATAIRE", "{{sig:prestataire}}"),
        ("CLIENT",      "{{sig:client}}"),
    ]

    for cell, header, (role, anchor) in zip(cells, headers, anchors):
        cell.width = Cm(8.5)
        # Title
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        # Filled fields
        for label in ("Nom et prénom", "Qualité", "Date"):
            pp = cell.add_paragraph()
            pp.paragraph_format.space_after = Pt(8)
            r = pp.add_run(f"{label} : ")
            r.font.size = Pt(10)
            r.bold = True
            placeholder = pp.add_run("____________________________________")
            placeholder.font.size = Pt(10)
            placeholder.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Signature zone
        pp = cell.add_paragraph()
        pp.paragraph_format.space_before = Pt(6)
        r = pp.add_run("Signature et cachet :")
        r.font.size = Pt(10)
        r.bold = True
        # Empty paragraphs to give the signature zone visible height
        for _ in range(4):
            cell.add_paragraph()
        # E-signature anchor (small, light, machine-detectable)
        pp = cell.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        anchor_run = pp.add_run(anchor)
        anchor_run.font.size = Pt(8)
        anchor_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        anchor_run.italic = True

        # Border around the signature box
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "8")
            border.set(qn("w:color"), "1F4E79")
            tcBorders.append(border)
        tcPr.append(tcBorders)

    # Mention obligatoire note below
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(
        "(Mention manuscrite « Lu et approuvé » avant signature de chaque Partie)"
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def build_one(src_path: Path) -> Path:
    raw = src_path.read_text(encoding="utf-8")

    # Strip trailing <!-- ... --> editor notes block
    raw = re.sub(r"<!--.*?-->", "", raw, flags=re.DOTALL).rstrip() + "\n"

    # Normalize line endings
    raw = raw.replace("\r\n", "\n")
    lines = raw.split("\n")

    # ─── Document setup ─────────────────────────────────────────────────
    doc = Document()

    # Default font Calibri 11
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # A4 + 1 inch margins (≈ 25.4 mm)
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(22)
    section.bottom_margin = Mm(22)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)

    add_page_number_footer(section)

    # ─── Walk lines ─────────────────────────────────────────────────────
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # Tables
        if is_table_row(line):
            rows, i = parse_table_block(lines, i)
            render_table(doc, rows)
            doc.add_paragraph()  # spacing after table
            continue

        # Horizontal rule
        if line.strip() == "---":
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(18)
            run = p.add_run(smart_quotes(text))
            run.font.name = "Calibri"
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            i += 1
            continue
        if line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(smart_quotes(text))
            run.font.name = "Calibri"
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            # Bottom border under article header
            p_pr = p._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "2")
            bottom.set(qn("w:color"), "1F4E79")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
            i += 1

            # Special-case: replace the markdown SIGNATURES section with a
            # rich visual signature block (boxes + anchor markers for e-sig).
            if text.upper().startswith("SIGNATURES"):
                render_signature_block(doc)
                # Skip ahead until the next H1/H2 heading or end of doc.
                while i < len(lines):
                    nxt = lines[i].rstrip()
                    if nxt.startswith("# ") or nxt.startswith("## "):
                        break
                    i += 1
            continue
        if line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(smart_quotes(text))
            run.font.name = "Calibri"
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            i += 1
            continue

        # Bullet list (- ...)
        if line.lstrip().startswith("- "):
            indent = (len(line) - len(line.lstrip())) // 2
            text = line.lstrip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.6 + 0.6 * indent)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, text)
            i += 1
            continue

        # Numbered list (1. ... or 1) ...)
        m = re.match(r"^(\s*)(\d+)\. (.+)$", line)
        if m:
            text = m.group(3)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, text)
            i += 1
            continue

        # Regular paragraph: collapse following non-empty lines that are not
        # headings/lists/tables/hr into one logical paragraph.
        buf = [line]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].rstrip()
            if not nxt.strip():
                break
            if nxt.startswith("#") or nxt.startswith("- ") or is_table_row(nxt) or nxt.strip() == "---":
                break
            if re.match(r"^\s*\d+\. ", nxt):
                break
            buf.append(nxt)
            j += 1
        para_text = " ".join(buf).strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, para_text)
        i = j

    dst = src_path.with_suffix(".docx")
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst)
    return dst


def main():
    if len(sys.argv) >= 2:
        sources = [Path(sys.argv[1]).resolve()]
    else:
        sources = DEFAULT_SOURCES
    for src in sources:
        if not src.exists():
            print(f"SKIP: missing {src}")
            continue
        out = build_one(src)
        print(f"OK -> {out}  ({out.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
